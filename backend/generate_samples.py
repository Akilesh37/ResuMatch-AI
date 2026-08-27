import os
from pathlib import Path
from docx import Document
from PIL import Image, ImageDraw, ImageFont

SAMPLES_DIR = Path(__file__).resolve().parent.parent / "samples"
SAMPLES_DIR.mkdir(parents=True, exist_ok=True)

# 1. Text Resume: Senior AI / Machine Learning Engineer
AI_RESUME_TEXT = """DR. ELENA ROSTOVA
Email: elena.rostova.ai@gmail.com | Phone: +1-650-555-0144 | Palo Alto, CA

PROFESSIONAL SUMMARY
Senior Machine Learning & AI Research Engineer with 5 years of industry experience developing transformer architectures, large language model (LLM) fine-tuning pipelines, and high-throughput vector retrieval systems.

WORK EXPERIENCE
Lead AI Engineer | DeepIntelligence Labs (2022 - Present)
- Architected enterprise RAG system with PyTorch, HuggingFace Transformers, and Vector Databases (FAISS/ChromaDB).
- Deployed real-time inference microservices with FastAPI and Docker on AWS EKS with Kubernetes.
- Optimized semantic embedding computation, reducing latency by 40%.

Machine Learning Engineer | Apex Data Systems (2019 - 2022)
- Built NLP classification models, entity extraction pipelines (NER), and sentiment analysis engines.
- Collaborated in cross-functional agile teams using CI/CD and Git.

TECHNICAL SKILLS
Python, PyTorch, Transformers, Deep Learning, Machine Learning, LLMs, LangChain, Vector Databases, FastAPI, Docker, Kubernetes, AWS, SQL, Scikit-Learn, Git, NLP.

EDUCATION
Ph.D. in Computer Science (Artificial Intelligence) | Stanford University (2019)
B.S. in Computer Science | UC Berkeley (2015)
"""

# 2. DOCX Resume: Full Stack Developer
def generate_docx_resume():
    doc = Document()
    doc.add_heading("DAVID MILLER", level=0)
    p_contact = doc.add_paragraph("Email: david.miller.tech@yahoo.com | Phone: +1-312-555-9012 | Chicago, IL")
    
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph("High-performing Full Stack Developer with 4 years of experience building modern web applications with React, TypeScript, Node.js, and PostgreSQL.")
    
    doc.add_heading("Work Experience", level=1)
    doc.add_paragraph("Senior Full Stack Developer | NextGen Web Solutions (2021 - Present)\n- Developed responsive frontend applications with React, Next.js, and TypeScript.\n- Built RESTful APIs and microservices in Node.js and Python with PostgreSQL databases.\n- Implemented Redis caching layer and containerized deployments with Docker.")
    
    doc.add_paragraph("Frontend Developer | ByteWave Media (2019 - 2021)\n- Developed UI components using JavaScript, React, HTML5, CSS3, and Tailwind CSS.\n- Participated in Agile sprint planning and Git version control.")
    
    doc.add_heading("Technical Skills", level=1)
    doc.add_paragraph("JavaScript, TypeScript, React, Next.js, Node.js, Python, PostgreSQL, REST API, Redis, Docker, Tailwind CSS, Git, HTML5, CSS3, Agile.")
    
    doc.add_heading("Education", level=1)
    doc.add_paragraph("B.S. in Computer Science | University of Illinois Urbana-Champaign (2019)")
    
    docx_path = SAMPLES_DIR / "sample_fullstack_dev.docx"
    doc.save(docx_path)
    print(f"Generated DOCX sample: {docx_path}")

# 3. Text Resume: DevOps & Infrastructure Engineer
DEVOPS_RESUME_TEXT = """SARAH JENNINGS
Email: sarah.jennings.devops@cloudmesh.io | Phone: +1-512-555-3389 | Austin, TX

SUMMARY
DevOps & Cloud Infrastructure Engineer with 3.5 years of experience automating CI/CD pipelines, Kubernetes orchestrations, and Terraform infrastructure-as-code across AWS and GCP.

PROFESSIONAL EXPERIENCE
DevOps Engineer | CloudScale Systems (2021 - Present)
- Provisioned infrastructure using Terraform and Ansible on AWS (EC2, S3, RDS, EKS).
- Configured Kubernetes clusters, Helm charts, and Docker containerization.
- Automated CI/CD pipelines using GitHub Actions, ensuring zero-downtime deployments.

Systems Administrator | NetCore Solutions (2019 - 2021)
- Managed Linux server administration (Ubuntu, RedHat), Bash automation scripting, and Nginx reverse proxies.

SKILLS
Docker, Kubernetes, AWS, Terraform, CI/CD, Linux, Bash, Git, Ansible, GCP, Python, Nginx, Security.

EDUCATION
B.S. in Information Technology | University of Texas at Austin (2019)
"""

# 4. Scanned PNG Resume (for Tesseract OCR testing)
def generate_scanned_image_resume():
    width, height = 1200, 1500
    image = Image.new("RGB", (width, height), color=(255, 255, 255))
    draw = ImageDraw.Draw(image)
    
    lines = [
        ("MICHAEL R. CHANG", 32, True),
        ("Email: michael.chang.ocr@gmail.com | Phone: 415-555-8900", 18, False),
        ("", 10, False),
        ("PROFESSIONAL SUMMARY", 22, True),
        ("Data Scientist & Machine Learning Specialist with 3 years of experience building", 18, False),
        ("predictive models, data pipelines, and analytical dashboards.", 18, False),
        ("", 10, False),
        ("WORK EXPERIENCE", 22, True),
        ("Data Scientist - QuantEdge Analytics (2021 - Present)", 19, True),
        ("- Engineered machine learning algorithms with Python, Scikit-Learn, and Pandas.", 18, False),
        ("- Built executive BI dashboards using Tableau, Power BI, and SQL.", 18, False),
        ("- Processed large datasets using Spark and SQL database queries.", 18, False),
        ("", 10, False),
        ("TECHNICAL SKILLS", 22, True),
        ("Python, SQL, Machine Learning, Data Science, Data Analysis, Pandas, NumPy,", 18, False),
        ("Scikit-Learn, Tableau, Power BI, Spark, Git, Deep Learning.", 18, False),
        ("", 10, False),
        ("EDUCATION", 22, True),
        ("Master of Science in Data Science | New York University (2021)", 18, False),
        ("Bachelor of Science in Statistics (2019)", 18, False),
    ]
    
    y = 80
    for text, size, is_bold in lines:
        if not text:
            y += 20
            continue
        draw.text((100, y), text, fill=(20, 25, 35))
        y += size + 16
        
    png_path = SAMPLES_DIR / "sample_scanned_resume.png"
    image.save(png_path)
    print(f"Generated OCR image sample: {png_path}")

def main():
    # Save TXT samples
    with open(SAMPLES_DIR / "sample_ai_engineer.txt", "w") as f:
        f.write(AI_RESUME_TEXT)
    with open(SAMPLES_DIR / "sample_devops_engineer.txt", "w") as f:
        f.write(DEVOPS_RESUME_TEXT)
        
    generate_docx_resume()
    generate_scanned_image_resume()
    print("All sample resumes created successfully in samples/ directory.")

if __name__ == "__main__":
    main()
