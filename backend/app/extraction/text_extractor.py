import io
import os
import logging
from pathlib import Path
from typing import Dict, Any
import pdfplumber
from docx import Document
from pypdf import PdfReader

from .ocr_engine import extract_text_from_image, extract_text_from_pdf_ocr

logger = logging.getLogger(__name__)

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp", ".gif"}
DOCX_EXTENSIONS = {".docx", ".doc"}
PDF_EXTENSIONS = {".pdf"}
TXT_EXTENSIONS = {".txt", ".rtf", ".md"}

def extract_text_from_pdf(file_bytes: bytes) -> Dict[str, Any]:
    """
    Extract text natively from PDF, falling back to OCR if scanned or empty.
    """
    native_text = ""
    page_count = 0
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    native_text += text + "\n"
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}. Trying pypdf...")
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            page_count = len(reader.pages)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    native_text += text + "\n"
        except Exception as e2:
            logger.warning(f"pypdf also failed: {e2}")

    clean_native = native_text.strip()
    
    # If native text is too short (< 30 chars), treat as scanned PDF and invoke OCR
    if len(clean_native) < 30:
        logger.info("PDF contains negligible native text. Falling back to OCR pipeline...")
        ocr_result = extract_text_from_pdf_ocr(file_bytes)
        if ocr_result.get("success") and len(ocr_result.get("text", "")) > 10:
            return {
                "success": True,
                "text": ocr_result["text"],
                "is_ocr": True,
                "ocr_confidence": ocr_result.get("confidence", 80.0),
                "page_count": page_count,
                "extraction_method": "tesseract_ocr_scanned_pdf"
            }
        # If OCR also failed or returned minimal text, return whatever native text we had
        return {
            "success": bool(clean_native),
            "text": clean_native or ocr_result.get("text", ""),
            "is_ocr": ocr_result.get("is_ocr", False),
            "ocr_confidence": ocr_result.get("confidence", 0.0),
            "page_count": page_count,
            "extraction_method": "native_pdf_partial"
        }

    return {
        "success": True,
        "text": clean_native,
        "is_ocr": False,
        "ocr_confidence": None,
        "page_count": page_count,
        "extraction_method": "native_pdf"
    }

def extract_text_from_docx(file_bytes: bytes) -> Dict[str, Any]:
    """
    Extract paragraphs and table text from DOCX files.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract table text
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    tables_text.append(" | ".join(row_text))
                    
        combined = "\n".join(paragraphs + tables_text).strip()
        return {
            "success": True,
            "text": combined,
            "is_ocr": False,
            "ocr_confidence": None,
            "page_count": None,
            "extraction_method": "python_docx"
        }
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return {
            "success": False,
            "text": "",
            "is_ocr": False,
            "ocr_confidence": None,
            "error": str(e),
            "extraction_method": "python_docx_failed"
        }

def extract_text_from_txt(file_bytes: bytes) -> Dict[str, Any]:
    """
    Extract text from plaintext or markdown files with encoding detection.
    """
    for encoding in ["utf-8", "latin-1", "cp1252", "iso-8859-1"]:
        try:
            text = file_bytes.decode(encoding).strip()
            return {
                "success": True,
                "text": text,
                "is_ocr": False,
                "ocr_confidence": None,
                "page_count": 1,
                "extraction_method": f"plaintext_{encoding}"
            }
        except UnicodeDecodeError:
            continue
            
    return {
        "success": False,
        "text": "",
        "is_ocr": False,
        "ocr_confidence": None,
        "error": "Failed to decode text file with standard encodings",
        "extraction_method": "plaintext_failed"
    }

def extract_resume_content(file_bytes: bytes, filename: str, content_type: str = "") -> Dict[str, Any]:
    """
    Universal extractor for any supported resume format (PDF, DOCX, TXT, Images).
    """
    ext = Path(filename).suffix.lower()
    
    if ext in PDF_EXTENSIONS or "pdf" in content_type.lower():
        result = extract_text_from_pdf(file_bytes)
    elif ext in DOCX_EXTENSIONS or "word" in content_type.lower() or "officedocument" in content_type.lower():
        result = extract_text_from_docx(file_bytes)
    elif ext in IMAGE_EXTENSIONS or "image/" in content_type.lower():
        ocr_res = extract_text_from_image(file_bytes)
        result = {
            "success": ocr_res["success"],
            "text": ocr_res["text"],
            "is_ocr": True,
            "ocr_confidence": ocr_res["confidence"],
            "page_count": 1,
            "extraction_method": "tesseract_ocr_image"
        }
    elif ext in TXT_EXTENSIONS or "text/" in content_type.lower():
        result = extract_text_from_txt(file_bytes)
    else:
        # Try generic plaintext fallback, if fails try OCR as image
        try:
            result = extract_text_from_txt(file_bytes)
            if not result.get("success") or len(result.get("text", "")) < 20:
                result = extract_text_from_pdf(file_bytes)
        except Exception:
            result = {
                "success": False,
                "text": "",
                "is_ocr": False,
                "ocr_confidence": None,
                "error": f"Unsupported file extension: {ext}"
            }
            
    # Add metadata
    if result.get("success"):
        text = result.get("text", "")
        result["char_count"] = len(text)
        result["word_count"] = len(text.split())
        
    return result
